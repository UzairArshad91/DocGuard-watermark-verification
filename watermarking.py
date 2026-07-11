import os
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import DictionaryObject, NameObject, TextStringObject, ArrayObject, NumberObject
import openpyxl
# optional image drawing imports are loaded when needed for XLSX
from reportlab.pdfgen import canvas
from io import BytesIO
import hashlib
from openpyxl.comments import Comment

def get_sig_hash(filepath, ext):
    if ext == ".docx":
        d = Document(filepath)
        sig = d.core_properties.comments
        visible = "\n".join(
            paragraph.text
            for section in d.sections
            for paragraph in section.footer.paragraphs
            if paragraph.text
        )
    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(filepath)
        sig = wb.properties.description
        # read visible watermark from A1 comment (do not touch cell value)
        try:
            ws = wb.active
            comment = ws["A1"].comment
            visible = comment.text if comment is not None else ""
        except Exception:
            visible = ""
    elif ext == ".pdf":
        # read signature metadata and collect any visible watermark annotations marked with /NM == 'WATERMARK'
        reader = PdfReader(filepath)
        sig = reader.metadata.get("/Signature")
        visible_parts = []
        try:
            for page in reader.pages:
                annots = page.get("/Annots")
                if not annots:
                    continue
                for a in annots:
                    try:
                        obj = a.get_object()
                        nm = obj.get("/NM")
                        if nm == TextStringObject("WATERMARK"):
                            c = obj.get("/Contents")
                            if c:
                                visible_parts.append(str(c))
                    except Exception:
                        continue
        except Exception:
            pass
        visible = "\n".join(visible_parts)
    else:
        sig = visible = None
    combined = (sig or "") + (visible or "")
    return hashlib.sha256(combined.encode()).hexdigest()

def add_watermark(filepath, name, email, user_id, sig):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        doc = Document(filepath)
        doc.core_properties.author = name
        footer = doc.sections[0].footer
        if footer.paragraphs:
            footer.paragraphs[0].text = f"Sent by: {name} ({email})"
            for paragraph in footer.paragraphs[1:]:
                paragraph.text = ""
        else:
            footer.add_paragraph(f"Sent by: {name} ({email})")
        doc.core_properties.comments = sig
        doc.save(filepath)
    
    elif ext == ".pdf":
        reader = PdfReader(filepath)
        writer = PdfWriter()
        # remove existing watermark annotations (with /NM == 'WATERMARK') and add new FreeText annotation at bottom
        for page in reader.pages:
            annots = page.get("/Annots")
            new_annots = []
            if annots:
                for a in annots:
                    try:
                        obj = a.get_object()
                        if obj.get("/NM") == TextStringObject("WATERMARK"):
                            continue
                        new_annots.append(a)
                    except Exception:
                        new_annots.append(a)
            # create new watermark annotation
            try:
                media = page.mediabox
                width = float(media.width)
                # place small watermark at bottom center
                llx = width * 0.1
                lly = 20
                urx = width * 0.9
                ury = 40
                annot = DictionaryObject()
                annot.update({
                    NameObject("/Type"): NameObject("/Annot"),
                    NameObject("/Subtype"): NameObject("/FreeText"),
                    NameObject("/Rect"): ArrayObject([NumberObject(llx), NumberObject(lly), NumberObject(urx), NumberObject(ury)]),
                    NameObject("/Contents"): TextStringObject(f"Sent by: {name} ({email})"),
                    NameObject("/NM"): TextStringObject("WATERMARK"),
                    NameObject("/F"): NumberObject(4),
                })
                ref = writer._add_object(annot)
                new_annots.append(ref)
                if new_annots:
                    page[NameObject("/Annots")] = ArrayObject(new_annots)
            except Exception:
                pass
            writer.add_page(page)

        writer.add_metadata({"/Author": name, "/Signature": sig})
        with open(filepath, "wb") as f:
            writer.write(f)

    elif ext in (".xlsx", ".xls"):
        wb = openpyxl.load_workbook(filepath)
        wb.properties.creator = name
        ws = wb.active
        watermark_text = f"Sent by: {name} ({email})"
        # write visible watermark into A1 cell comment (leave cell value intact)
        try:
            try:
                comment = Comment(watermark_text, name)
                ws["A1"].comment = comment
            except Exception:
                # if comment API not available or fails, fall back to header without touching A1 value
                try:
                    ws.oddHeader.center = watermark_text
                except Exception:
                    pass
        except Exception:
            pass

        wb.properties.description = sig
        wb.save(filepath)